# make_sample_template.py
"""
샘플 한국어 보안 사고 보고서 템플릿(.docx)을 생성한다.
사용자가 자신의 양식이 없을 때 이걸 그대로 쓰거나, 형식만 참고할 수 있다.

사용자 회사 양식이 따로 있다면 이 스크립트를 실행할 필요가 없다.
template_ko.docx 를 직접 작성하고, 안에 {{title}} {{date}} 같은
placeholder만 박아두면 된다.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


def build(output_path: str = "template_ko.docx") -> None:
    doc = Document()

    # 기본 폰트
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    # 표지 제목
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("보안 사고 분석 보고서")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    # 표지 메타 표
    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Light Grid Accent 1"
    meta.autofit = True

    def _row(i, k, v):
        meta.cell(i, 0).text = k
        meta.cell(i, 1).text = v
        for cell in meta.rows[i].cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    _row(0, "사고 ID",   "{{incident_id}}")
    _row(1, "사고 제목", "{{title}}")
    _row(2, "발생일",    "{{date}}")
    _row(3, "심각도",    "{{severity}}")

    doc.add_paragraph()  # spacer

    # 1. 사고 개요
    h1 = doc.add_heading("1. 사고 개요", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    doc.add_paragraph("{{summary}}")

    doc.add_heading("2. 영향 자산", level=1)
    doc.add_paragraph("{{assets}}")

    doc.add_heading("3. MITRE ATT&CK 매핑", level=1)
    doc.add_paragraph("관련 기법 요약: {{techniques_short}}")
    doc.add_paragraph("상세 내역:")
    doc.add_paragraph("{{techniques}}")

    doc.add_heading("4. 침해 지표 (IOC)", level=1)
    doc.add_paragraph("{{iocs}}")

    doc.add_heading("5. 타임라인", level=1)
    doc.add_paragraph("{{timeline}}")

    doc.add_heading("6. 근본 원인 분석", level=1)
    doc.add_paragraph("{{root_cause}}")

    doc.add_heading("7. 권고사항", level=1)
    doc.add_paragraph("{{recommendations}}")

    # 바닥글에도 placeholder 잘 동작하는지 데모용으로 한 줄
    section = doc.sections[0]
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.text = "사고 ID: {{incident_id}} · 발생일: {{date}}"

    out = Path(output_path)
    doc.save(str(out))
    print(f"생성됨: {out.resolve()}")


if __name__ == "__main__":
    import sys
    build(sys.argv[1] if len(sys.argv) > 1 else "template_ko.docx")

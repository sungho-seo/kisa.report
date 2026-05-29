# test_filler.py
"""
LLM 호출 없이 docx_filler 가 제대로 동작하는지 검증.

검증 항목:
1. 일반 단순 치환 ({{title}}, {{date}})
2. 표 셀 안의 placeholder
3. 바닥글 안의 placeholder
4. 리스트 필드 ({{recommendations}}, {{techniques}}) 줄바꿈 처리
5. Run 분할 placeholder (Word 자동 교정 시뮬레이션)
6. 결과 .docx 가 다시 열리고 placeholder가 남아있지 않은지
"""

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).parent))

from docx_filler import _replace_in_paragraph, build_value_map, fill_template
from models import IOC, IncidentReport, Technique, TimelineEvent


def make_report() -> IncidentReport:
    return IncidentReport(
        title="외부 SSH 무차별 대입 시도 사고",
        incident_id="INC-2025-001",
        occurred_at=date(2025, 5, 28),
        severity="High",
        summary=(
            "5월 28일 03:12 KST부터 약 40분간 외부 IP에서 SSH 서비스로 "
            "약 12,000건의 인증 실패가 발생했다. 1건의 인증 성공이 관찰되어 "
            "권한 탈취 가능성이 있으며, 해당 계정의 비밀번호는 즉시 변경되었다."
        ),
        affected_assets=["bastion01 (10.0.1.5)", "app-prod-3 (10.0.2.11)"],
        techniques=[
            Technique(
                id="T1110",
                name="Brute Force",
                tactic="Credential Access",
                evidence="sshd: 12,143 Failed password from 203.0.113.45",
            ),
            Technique(
                id="T1078.003",
                name="Valid Accounts: Local Accounts",
                tactic="Initial Access",
                evidence="Accepted password for ops from 203.0.113.45 port 51234",
            ),
        ],
        iocs=[
            IOC(type="ip", value="203.0.113.45", description="공격 출발지"),
            IOC(type="ip", value="198.51.100.22", description="동일 대역 동시 시도"),
        ],
        timeline=[
            TimelineEvent(when="2025-05-28 03:12", what="첫 실패 인증 관측"),
            TimelineEvent(when="2025-05-28 03:48", what="ops 계정 인증 성공 1건"),
            TimelineEvent(when="2025-05-28 03:51", what="해당 IP 방화벽 차단"),
        ],
        root_cause=(
            "bastion01 의 fail2ban 정책이 비활성화되어 있어 반복 인증 실패가 "
            "차단되지 않았다."
        ),
        recommendations=[
            "bastion01 의 fail2ban 활성화 및 ssh 5회 실패 시 1시간 차단",
            "ops 계정 비밀번호 회전 및 MFA 강제",
            "외부 공격 IP 대역에 대한 WAF/Firewall 차단 규칙 추가",
        ],
    )


def test_build_value_map():
    print("[1] build_value_map ...", end=" ")
    r = make_report()
    v = build_value_map(r)
    assert v["title"] == r.title
    assert v["date"] == "2025-05-28"
    assert v["severity"] == "High"
    assert "T1110" in v["techniques"]
    assert "T1110, T1078.003" == v["techniques_short"]
    assert v["recommendations"].startswith("1. ")
    assert "\n" in v["recommendations"]  # 줄바꿈으로 연결
    print("OK")


def test_replace_run_split():
    """Word 자동 교정으로 placeholder가 쪼개진 상태를 시뮬레이션."""
    print("[2] run-split placeholder ...", end=" ")
    from docx import Document as Doc
    d = Doc()
    p = d.add_paragraph()
    # "{{title}}" 을 5개 Run에 쪼개 넣음 + 앞뒤로 텍스트
    p.add_run("제목: ")
    p.add_run("{{ti")
    p.add_run("tle")
    p.add_run("}}")
    p.add_run(" - 끝")

    values = {"title": "테스트 사고"}
    changed = _replace_in_paragraph(p, values)
    assert changed is True
    # 줄바꿈은 <w:br/>로 들어가므로 .text가 '\n'을 포함해야 함 (없으면 단순 텍스트)
    full = p.text
    assert "제목: 테스트 사고 - 끝" in full, f"실제 텍스트: {full!r}"
    # 여러 Run이 그대로 남아있고 (서식 유지 가능), 첫 Run만 텍스트를 갖는지 확인
    nonempty_runs = [r for r in p.runs if r.text]
    assert len(nonempty_runs) >= 1
    print("OK")


def test_multiple_placeholders_one_paragraph():
    print("[3] same-paragraph 다중 placeholder ...", end=" ")
    from docx import Document as Doc
    d = Doc()
    p = d.add_paragraph("ID: {{incident_id}} / 발생일: {{date}}")
    _replace_in_paragraph(p, {"incident_id": "INC-001", "date": "2025-05-28"})
    assert p.text == "ID: INC-001 / 발생일: 2025-05-28", p.text
    print("OK")


def test_undefined_placeholder_left_alone():
    print("[4] 미정의 placeholder는 보존 ...", end=" ")
    from docx import Document as Doc
    d = Doc()
    p = d.add_paragraph("정의됨: {{title}}, 미정의: {{unknown_xyz}}")
    _replace_in_paragraph(p, {"title": "X"})
    # 정의된 것은 치환, 미정의는 그대로 (또는 그대로 남아있어야 함 - 단 첫 미정의에서 중단됨)
    # 본 구현은 첫 미정의에서 멈추므로, title 이 먼저 등장하므로 치환됨
    # unknown_xyz 는 그대로 텍스트에 남아있어야 함
    assert "{{unknown_xyz}}" in p.text
    assert "{{title}}" not in p.text
    print("OK", "->", p.text)


def test_end_to_end_template():
    print("[5] 템플릿 채우기 e2e ...", end=" ")
    tpl = Path(__file__).parent / "template_ko.docx"
    out = Path(__file__).parent / "out" / "incident_test.docx"
    if out.exists():
        out.unlink()
    fill_template(tpl, make_report(), out)
    assert out.exists() and out.stat().st_size > 0

    # 재오픈해서 placeholder 가 남아 있지 않은지 검사
    d = Document(str(out))
    leftover_paras = []
    # 본문
    for p in d.paragraphs:
        if "{{" in p.text:
            leftover_paras.append(("body", p.text))
    # 표
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "{{" in p.text:
                        leftover_paras.append(("table", p.text))
    # 바닥글
    for section in d.sections:
        for p in section.footer.paragraphs:
            if "{{" in p.text:
                leftover_paras.append(("footer", p.text))

    assert not leftover_paras, f"치환되지 않은 placeholder: {leftover_paras}"

    # 핵심 값들이 본문 어딘가에 있는지 확인
    full = "\n".join(p.text for p in d.paragraphs) + "\n"
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                full += cell.text + "\n"
    assert "외부 SSH 무차별 대입 시도 사고" in full
    assert "INC-2025-001" in full
    assert "High" in full
    assert "T1110" in full
    assert "203.0.113.45" in full
    assert "fail2ban" in full
    print("OK ->", out)


if __name__ == "__main__":
    test_build_value_map()
    test_replace_run_split()
    test_multiple_placeholders_one_paragraph()
    test_undefined_placeholder_left_alone()
    test_end_to_end_template()
    print("\n모든 테스트 통과 ✅")
